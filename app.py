
# app.py
import streamlit as st
import sqlite3
import pandas as pd
import plotly.express as px
from datetime import datetime
from io import BytesIO
from docx import Document
import base64

st.set_page_config(page_title="Suivi des Opérations", layout="wide")

DB_FILE = "operations.db"

# Connexion et initialisation
@st.cache_resource
def get_connection():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    return conn

conn = get_connection()
cursor = conn.cursor()

def init_db():
    cursor.execute('''CREATE TABLE IF NOT EXISTS operations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nom TEXT, type TEXT, responsable TEXT, statut TEXT,
        avancement REAL DEFAULT 0, phases TEXT, date_creation TEXT
    )''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS journal (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operation_id INTEGER, action TEXT, auteur TEXT, date TEXT
    )''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS pieces_jointes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        operation_id INTEGER, nom TEXT, data_base64 TEXT
    )''')
    conn.commit()

init_db()

PHASES = [
    "Phase de montage", "Programmation", "Foncier", "Études",
    "DCE", "Attribution de marché", "Chantier", "Livraison",
    "Clôture technique", "Clôture financière"
]
STATUTS = ["🟡 À l’étude", "🟢 En cours", "🔴 Bloqué", "✅ Clôturé"]

# Fonctions utilitaires
def ajouter_operation(nom, type_op, responsable, statut):
    date = datetime.now().isoformat()
    cursor.execute("INSERT INTO operations (nom, type, responsable, statut, phases, date_creation) VALUES (?, ?, ?, ?, ?, ?)",
                   (nom, type_op, responsable, statut, "[]", date))
    conn.commit()

def get_operations(filtres=None):
    df = pd.read_sql("SELECT * FROM operations", conn)
    return df

def get_journal(operation_id):
    return pd.read_sql(f"SELECT * FROM journal WHERE operation_id={operation_id} ORDER BY date DESC", conn)

def get_pieces(operation_id):
    return pd.read_sql(f"SELECT * FROM pieces_jointes WHERE operation_id={operation_id}", conn)

# --- Interface
menu = st.sidebar.radio("📂 Navigation", ["Vue Opérations", "Vue Détails", "Vue Manager", "🔄 Réinitialiser DB"])

# --- Vue Opérations ---
if menu == "Vue Opérations":
    st.title("📋 Suivi des opérations")
    with st.expander("➕ Nouvelle opération"):
        nom = st.text_input("Nom")
        type_op = st.selectbox("Type", ["OPP", "VEFA", "AMO", "Mandat"])
        responsable = st.text_input("Responsable")
        statut = st.selectbox("Statut", STATUTS)
        if st.button("Ajouter"):
            ajouter_operation(nom, type_op, responsable, statut)
            st.success("Opération ajoutée")

    st.subheader("🔍 Liste et filtres")
    df = get_operations()
    col1, col2, col3 = st.columns(3)
    f_type = col1.selectbox("Filtrer Type", ["Tous"] + df["type"].unique().tolist())
    f_statut = col2.selectbox("Filtrer Statut", ["Tous"] + STATUTS)
    f_resp = col3.selectbox("Filtrer Responsable", ["Tous"] + df["responsable"].unique().tolist())

    if f_type != "Tous":
        df = df[df["type"] == f_type]
    if f_statut != "Tous":
        df = df[df["statut"] == f_statut]
    if f_resp != "Tous":
        df = df[df["responsable"] == f_resp]

    st.dataframe(df)

# --- Vue Détails ---
elif menu == "Vue Détails":
    st.title("📁 Détail d'une opération")
    df = get_operations()
    op_nom = st.selectbox("Choisir une opération", df["nom"] if not df.empty else [])
    if op_nom:
        op = df[df["nom"] == op_nom].iloc[0]
        st.markdown(f"### 🔎 {op_nom}")
        st.write(f"**Type :** {op['type']}")
        st.write(f"**Responsable :** {op['responsable']}")
        st.write(f"**Statut :** {op['statut']}")
        st.write(f"**Créée le :** {op['date_creation'][:10]}")

        st.subheader("📌 Phases & Avancement")
        phase_check = st.multiselect("Cochez les phases terminées :", PHASES)
        avancement = round(len(phase_check) / len(PHASES) * 100, 1)
        st.progress(avancement / 100)
        st.write(f"**Avancement :** {avancement} %")

        st.subheader("🗒️ Journal des actions")
        action = st.text_input("Nouvelle entrée")
        auteur = st.text_input("Auteur", value="Système")
        if st.button("Ajouter au journal"):
            cursor.execute("INSERT INTO journal (operation_id, action, auteur, date) VALUES (?, ?, ?, ?)",
                           (op["id"], action, auteur, datetime.now().isoformat()))
            conn.commit()
            st.success("Ajouté")
        st.dataframe(get_journal(op["id"]))

        st.subheader("📎 Pièces jointes")
        upload = st.file_uploader("Téléverser un fichier", type=["pdf", "jpg", "png"])
        if upload:
            data = base64.b64encode(upload.read()).decode()
            cursor.execute("INSERT INTO pieces_jointes (operation_id, nom, data_base64) VALUES (?, ?, ?)",
                           (op["id"], upload.name, data))
            conn.commit()
            st.success("Fichier ajouté")
        files = get_pieces(op["id"])
        for _, row in files.iterrows():
            st.download_button(f"⬇️ {row['nom']}", base64.b64decode(row["data_base64"]), file_name=row["nom"])

        st.subheader("📄 Export Word")
        def export_word(op):
            doc = Document()
            doc.add_heading(op['nom'], 0)
            doc.add_paragraph(f"Type : {op['type']}")
            doc.add_paragraph(f"Responsable : {op['responsable']}")
            doc.add_paragraph(f"Statut : {op['statut']}")
            doc.add_paragraph("Phases : " + ", ".join(phase_check))
            doc.add_paragraph(f"Avancement : {avancement} %")
            doc.add_heading("Journal", level=1)
            for _, row in get_journal(op["id"]).iterrows():
                doc.add_paragraph(f"{row['date'][:10]} - {row['auteur']} : {row['action']}")
            b = BytesIO()
            doc.save(b)
            b.seek(0)
            return b
        if st.button("📤 Générer rapport Word"):
            st.download_button("Télécharger", export_word(op), file_name="rapport.docx")

# --- Vue Manager ---
elif menu == "Vue Manager":
    st.title("📊 Vue Manager")
    df = get_operations()
    st.metric("Nombre total", len(df))
    st.metric("En cours", len(df[df["statut"] == "🟢 En cours"]))
    st.metric("Clôturées", len(df[df["statut"] == "✅ Clôturé"]))
    st.metric("Bloquées", len(df[df["statut"] == "🔴 Bloqué"]))
    fig = px.histogram(df, x="statut", color="responsable", title="Statuts par responsable")
    st.plotly_chart(fig)

# --- Reset DB ---
elif menu == "🔄 Réinitialiser DB":
    if st.button("⚠️ Réinitialiser toute la base de données"):
        cursor.execute("DROP TABLE IF EXISTS operations")
        cursor.execute("DROP TABLE IF EXISTS journal")
        cursor.execute("DROP TABLE IF EXISTS pieces_jointes")
        conn.commit()
        init_db()
        st.success("Base réinitialisée")